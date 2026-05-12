"""Add native Word comments (.docx) to findings.

Manipulates python-docx underlying OOXML to insert w:comment elements
and document-level commentRangeStart/commentRangeEnd markers.
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


@dataclass(frozen=True)
class Finding:
    file: str
    page: int
    target: str
    category: str
    suggestion: str
    severity: str = "medium"
    source: str = "rule"
    occurrence: int = 0
    fallback_rect: tuple[float, float, float, float] | None = None
    end_of_doc: bool = False


def _extract_display_target(target: str) -> str:
    """Strip auto: prefixes so we can search for real text."""
    if target.startswith("auto:"):
        parts = target.split(":")
        if len(parts) >= 3:
            # e.g. auto:stat-style:p:5:120 -> try to find a real snippet
            return ":".join(parts[2:])
    return target


def _ensure_comments_part(document: Any) -> Any:
    """Get or create the word/comments.xml part."""
    # Search existing relationships
    for rel in document.part.rels.values():
        if "comments" in rel.reltype:
            return rel.target_part

    # Create new comments part using XmlPart with PackURI partname
    from docx.opc.part import XmlPart
    from docx.opc.packuri import PackURI

    comments_xml = (
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "</w:comments>"
    )
    element = parse_xml(comments_xml)
    part = XmlPart(
        partname=PackURI("/word/comments.xml"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        element=element,
        package=document.part.package,
    )
    document.part.package.parts.append(part)
    document.part.relate_to(part, RT.COMMENTS)
    return part


def _next_comment_id(comments_part: Any) -> int:
    """Return the next available comment id."""
    root = comments_part.element
    ids = [int(c.get(qn("w:id"), "0")) for c in root.findall(qn("w:comment"))]
    return max(ids, default=-1) + 1


def _add_comment_to_part(comments_part: Any, comment_id: int, text: str, author: str = "AI审读") -> None:
    """Append a w:comment element to the comments part."""
    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(comment_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    comment.set(qn("w:initials"), "AI")

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    comment.append(p)

    comments_part.element.append(comment)


def _find_run_with_text(paragraph: Any, target: str) -> Any | None:
    """Find the first run in a paragraph containing the target text."""
    for run in paragraph.runs:
        if target in run.text:
            return run
    return None


def _split_run(paragraph: Any, run: Any, target: str) -> Any | None:
    """Split a run so the target text is isolated in its own run.

    Returns the new run that contains only the target text, or None on failure.
    """
    text = run.text
    idx = text.find(target)
    if idx < 0:
        return None

    before = text[:idx]
    after = text[idx + len(target) :]

    rPr = run._r.find(qn("w:rPr"))
    p_elem = paragraph._p

    # Modify original run to hold 'before'
    t_elem = run._r.find(qn("w:t"))
    if t_elem is None:
        return None

    # If there is xml:space="preserve", preserve it
    space_attr = qn("xml:space")
    preserve = t_elem.get(space_attr)

    t_elem.text = before
    if before and before.endswith(" ") or before.startswith(" "):
        t_elem.set(space_attr, "preserve")
    elif preserve and not before:
        # Remove preserve if text becomes empty
        if space_attr in t_elem.attrib:
            del t_elem.attrib[space_attr]

    # Insert target run after original
    target_r = OxmlElement("w:r")
    if rPr is not None:
        target_r.append(deepcopy(rPr))
    target_t = OxmlElement("w:t")
    target_t.text = target
    if target.endswith(" ") or target.startswith(" "):
        target_t.set(space_attr, "preserve")
    target_r.append(target_t)

    # Insert after original run
    orig_index = list(p_elem).index(run._r)
    p_elem.insert(orig_index + 1, target_r)

    # Insert after run if needed
    if after:
        after_r = OxmlElement("w:r")
        if rPr is not None:
            after_r.append(deepcopy(rPr))
        after_t = OxmlElement("w:t")
        after_t.text = after
        if after.endswith(" ") or after.startswith(" "):
            after_t.set(space_attr, "preserve")
        after_r.append(after_t)
        p_elem.insert(orig_index + 2, after_r)

    return target_r


def _add_comment_markers(paragraph: Any, target_run: Any, comment_id: int) -> None:
    """Insert commentRangeStart before target_run and commentRangeEnd + commentReference after it."""
    p_elem = paragraph._p
    run_index = list(p_elem).index(target_run)

    # commentRangeStart
    crs = OxmlElement("w:commentRangeStart")
    crs.set(qn("w:id"), str(comment_id))
    p_elem.insert(run_index, crs)

    # commentRangeEnd (after target run)
    cre = OxmlElement("w:commentRangeEnd")
    cre.set(qn("w:id"), str(comment_id))
    p_elem.insert(run_index + 2, cre)

    # commentReference run
    ref_r = OxmlElement("w:r")
    ref_cr = OxmlElement("w:commentReference")
    ref_cr.set(qn("w:id"), str(comment_id))
    ref_r.append(ref_cr)
    p_elem.insert(run_index + 3, ref_r)


def _add_comment_to_paragraph(
    document: Any,
    paragraph: Any,
    target: str,
    comment_text: str,
    author: str = "AI审读",
) -> bool:
    """Add a single comment to a paragraph targeting specific text."""
    comments_part = _ensure_comments_part(document)
    comment_id = _next_comment_id(comments_part)

    run = _find_run_with_text(paragraph, target)
    if run is None:
        return False

    # If the run text is exactly the target, no need to split
    if run.text == target:
        target_run = run._r
    else:
        target_run = _split_run(paragraph, run, target)
        if target_run is None:
            return False

    _add_comment_markers(paragraph, target_run, comment_id)
    _add_comment_to_part(comments_part, comment_id, comment_text, author)
    return True


def _add_unlocated_summary(document: Any, findings: list[Finding]) -> int:
    """Append a summary paragraph for findings that could not be located."""
    if not findings:
        return 0

    body = document.element.body

    # Add blank line
    blank_p = OxmlElement("w:p")
    body.append(blank_p)

    title_p = OxmlElement("w:p")
    title_r = OxmlElement("w:r")
    title_rPr = OxmlElement("w:rPr")
    title_b = OxmlElement("w:b")
    title_rPr.append(title_b)
    title_r.append(title_rPr)
    title_t = OxmlElement("w:t")
    title_t.text = "未定位批注汇总（以下问题未能在原文中找到精确位置，请手动核查）"
    title_r.append(title_t)
    title_p.append(title_r)
    body.append(title_p)

    count = 0
    for f in findings:
        count += 1
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = f"{count}. [{f.category}] 段落{f.page} — {f.suggestion[:200]}"
        r.append(t)
        p.append(r)
        body.append(p)

    return count


def annotate_docx(docx_path: Path, findings: list[Finding], output_path: Path, same_name: bool = True) -> dict[str, Any]:
    """Add native Word comments for all findings and save annotated document."""
    result: dict[str, Any] = {
        "file": str(docx_path),
        "status": "skipped",
        "reason": "",
        "findings_configured": len(findings),
        "findings_annotated": 0,
        "missed": [],
        "output": "",
    }

    try:
        document = Document(str(docx_path))
    except Exception as exc:
        result["reason"] = f"open failed: {exc}"
        return result

    annotated = 0
    missed: list[dict[str, Any]] = []
    unlocated: list[Finding] = []

    try:
        for finding in findings:
            if Path(finding.file).name != docx_path.name:
                continue

            if finding.end_of_doc:
                unlocated.append(finding)
                continue

            target = _extract_display_target(finding.target)
            if not target.strip():
                missed.append({"finding": finding, "reason": "empty_target"})
                continue

            comment_text = f"[{finding.category}] {finding.suggestion}"
            success = False

            # Try the indicated paragraph first, then neighbours
            for attempt_para in (finding.page, finding.page - 1, finding.page + 1):
                if 0 <= attempt_para - 1 < len(document.paragraphs):
                    para = document.paragraphs[attempt_para - 1]
                    if target in para.text:
                        if _add_comment_to_paragraph(document, para, target, comment_text):
                            success = True
                            break

            if success:
                annotated += 1
            else:
                # Try all paragraphs as fallback
                for para in document.paragraphs:
                    if target in para.text:
                        if _add_comment_to_paragraph(document, para, target, comment_text):
                            success = True
                            break
                if success:
                    annotated += 1
                else:
                    missed.append({"finding": finding, "reason": "target_not_found"})
                    unlocated.append(finding)

        if unlocated:
            annotated += _add_unlocated_summary(document, unlocated)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        out_name = docx_path.name if same_name else f"{docx_path.stem}_annotated.docx"
        out_file = output_path.parent / out_name
        tmp_file = output_path.parent / f".{docx_path.stem}.tmp.docx"
        if tmp_file.exists():
            tmp_file.unlink()
        document.save(str(tmp_file))
        tmp_file.replace(out_file)

        result.update({
            "status": "done",
            "findings_annotated": annotated,
            "missed": missed,
            "output": str(out_file),
        })
    except Exception as exc:
        result["reason"] = f"annotate failed: {exc}"
        return result

    return result
