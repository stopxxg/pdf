"""Word-specific mechanical rule detectors.

Operates on python-docx Document objects, inspecting OOXML run properties
(italic, subscript, superscript) to flag formatting issues.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Any

from docx.oxml.ns import qn


@dataclass(frozen=True)
class Finding:
    file: str
    page: int
    target: str
    category: str
    suggestion: str
    severity: str = "medium"
    source: str = "rule"
    occurrence: int = 0
    fallback_rect: tuple[float, float, float, float] | None = None
    end_of_doc: bool = False


def _is_italic(run: Any) -> bool:
    if run.font.italic:
        return True
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return False
    return rPr.find(qn("w:i")) is not None


def _is_subscript(run: Any) -> bool:
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return False
    elem = rPr.find(qn("w:vertAlign"))
    return elem is not None and elem.get(qn("w:val")) == "subscript"


def _build_para_text_with_runs(paragraph: Any) -> tuple[str, list[tuple[int, int, Any]]]:
    """Return (full_text, [(start, end, run), ...]) for a paragraph."""
    full_text = ""
    runs: list[tuple[int, int, Any]] = []
    for run in paragraph.runs:
        text = run.text
        if text:
            start = len(full_text)
            full_text += text
            runs.append((start, len(full_text), run))
    return full_text, runs


def _find_run_at_offset(runs: list[tuple[int, int, Any]], offset: int) -> Any | None:
    for start, end, run in runs:
        if start <= offset < end:
            return run
    return None


def detect_word_stat_symbol_style(document: Any, filename: str) -> list[Finding]:
    """Detect statistical symbols that should be italic but are not."""
    findings: list[Finding] = []
    stat_patterns = [
        (re.compile(r"[pP]\s*(?:[<=>≤≥])\s*(?:0?\.\d+|\d+)"), "p", "p值中的p"),
        (re.compile(r"[I]\d+"), "I", "Moran's I中的I"),
        (re.compile(r"[R]\d+|R²"), "R", "相关系数R"),
        (re.compile(r"[F]\s*\(|F\d+"), "F", "F统计量"),
        (re.compile(r"[tzq]\s*(?:=|<|>|≥|≤|\()"), "tzq", "统计符号"),
    ]

    for para_idx, para in enumerate(document.paragraphs, start=1):
        full_text, runs = _build_para_text_with_runs(para)
        if not full_text.strip():
            continue

        for regex, symbol_label, desc in stat_patterns:
            for match in regex.finditer(full_text):
                symbol_offset = match.start()
                run = _find_run_at_offset(runs, symbol_offset)
                if run is None:
                    continue
                if not _is_italic(run):
                    target_text = match.group()
                    findings.append(
                        Finding(
                            filename,
                            para_idx,
                            target_text,
                            "公式/统计符号正斜体",
                            f"格式检查显示，此处{desc}为正体。建议排为斜体。",
                            severity="medium",
                            source="word-format",
                        )
                    )
    return findings


def detect_word_script_style(document: Any, filename: str) -> list[Finding]:
    """Detect missing subscripts for variable+digit patterns (e.g., I30, T1)."""
    findings: list[Finding] = []
    subscript_prefixes = set("ITRPXYZWVSDHCKMNpxyzwvsdhckmn")

    for para_idx, para in enumerate(document.paragraphs, start=1):
        full_text, runs = _build_para_text_with_runs(para)
        if not full_text.strip():
            continue

        i = 0
        while i < len(full_text) - 1:
            ch = full_text[i]
            if ch.isalpha() and ch in subscript_prefixes:
                letter_run = _find_run_at_offset(runs, i)
                if letter_run is None or not _is_italic(letter_run):
                    i += 1
                    continue

                j = i + 1
                digits: list[tuple[int, Any]] = []
                while j < len(full_text) and full_text[j].isdigit():
                    digit_run = _find_run_at_offset(runs, j)
                    if digit_run is not None:
                        digits.append((j, digit_run))
                    j += 1

                if 1 <= len(digits) <= 3:
                    digit_text = "".join(full_text[idx] for idx, _ in digits)
                    # Skip likely years (19xx, 20xx)
                    if len(digit_text) == 4 and digit_text.startswith(("19", "20")):
                        i = j
                        continue
                    # Skip if preceded by period or comma (likely citation)
                    if i > 0 and full_text[i - 1] in ".，,":
                        i = j
                        continue
                    # Skip if followed by closing bracket/parenthesis
                    if j < len(full_text) and full_text[j] in ")】）］]}":
                        i = j
                        continue
                    all_subscript = all(_is_subscript(run) for _, run in digits)
                    if not all_subscript:
                        target_text = ch + digit_text
                        findings.append(
                            Finding(
                                filename,
                                para_idx,
                                target_text,
                                "公式/下标格式",
                                f"格式检查显示，{target_text} 中的数字未显示为下标。建议将数字改为下标格式。",
                                severity="medium",
                                source="word-format",
                            )
                        )
                i = j
            else:
                i += 1
    return findings


def _compact(s: str) -> str:
    return re.sub(r"\s+", "", s)


def detect_word_text_rules(filename: str, para_idx: int, text: str) -> list[Finding]:
    """Pure text/regex based rules (mirrors PDF pipeline logic)."""
    findings: list[Finding] = []

    if para_idx == 1:
        front = text[:2400]
        if re.search(r"文献标识码：\s*(?:文章编号|中图分类号|\n)", front):
            findings.append(
                Finding(
                    filename, para_idx, "文献标识码", "期刊元数据",
                    "文献标识码疑似缺失。建议按期刊规范补充对应标识码。", "high",
                )
            )
        # Structured abstract check: support full-width, half-width, and plain text
        abstract_labels = ["目的", "方法", "结果", "结论"]
        label_patterns = [f"[{label}]" for label in abstract_labels]
        label_patterns += [f"［{label}］" for label in abstract_labels]
        label_patterns += [f"【{label}】" for label in abstract_labels]
        label_patterns += [f"{label}" for label in abstract_labels]
        found_labels = [p for p in label_patterns if p in front]
        review_keywords = ["综述", "进展", "展望", "review", "overview", "progress"]
        is_review = any(kw in front for kw in review_keywords)
        if not is_review and len(found_labels) < 3:
            missing_labels = [label for label in abstract_labels if not any(label in found for found in found_labels)]
            if missing_labels:
                findings.append(
                    Finding(
                        filename, para_idx, "摘", "摘要结构",
                        f"结构式摘要未检出{''.join(missing_labels)}。建议核查摘要栏目是否完整（若本文为综述可忽略）。", "medium",
                    )
                )

    text_targets: list[tuple[str, str, str]] = [
        ("http：", "文字/标点", "URL存在全角冒号，可能导致链接失效。建议改为半角“http://”或“https://”。"),
        ("https：", "文字/标点", "URL存在全角冒号，可能导致链接失效。建议改为半角“https://”。"),
        ("0. 001", "文字/标点", "数值“0. 001”中存在多余空格。建议改为“0.001”。"),
        ("，。", "文字/标点", "连续出现逗号和句号。建议删除多余标点。"),
        ("。。", "文字/标点", "连续出现两个句号。建议删除多余标点。"),
        ("..", "文字/标点", "连续出现两个英文句点。建议核查DOI、URL或参考文献标点。"),
        ("、、", "文字/标点", "连续出现两个顿号。建议删除多余顿号。"),
        ("本研仍", "文字/标点", "“本研”疑为“本研究”。建议补全。"),
        ("波段性", "文字/标点", "“波段性”在趋势描述中疑为“波动性”。建议核改。"),
        ("与和", "文字/标点", "“与和”连用不当。建议删除多余连接词。"),
        ("摘 要", "文字/标点", "“摘 要”中间有多余空格，应改为“摘要”。"),
    ]
    for target, category, suggestion in text_targets:
        occurrence = 0
        start = 0
        while True:
            hit = text.find(target, start)
            if hit < 0:
                break
            findings.append(
                Finding(
                    filename, para_idx, target, category, suggestion,
                    occurrence=occurrence,
                )
            )
            occurrence += 1
            start = hit + len(target)

    # Regex-based rules
    regex_rules: list[tuple[str, str, str]] = [
        (r"[pP]\s*<\s*0\.\s+0[15]", "公式/统计表达", "p值表达存在多余空格或断裂风险。建议统一为紧凑形式，并核查p是否斜体。"),
        (r"0\.\s+\d+", "文字/标点", "小数点后存在多余空格，建议删除空格。"),
        (r"图\s+\d+", "文字/标点", "“图”与编号之间存在多余空格，建议改为“图1”格式。"),
        (r"表\s+\d+", "文字/标点", "“表”与编号之间存在多余空格，建议改为“表1”格式。"),
        (r"et al\.[A-Z]", "文字/标点", "“et al.”后缺少空格，建议改为“et al. Author”。"),
    ]
    for regex, category, suggestion in regex_rules:
        occurrence = 0
        for match in re.finditer(regex, text):
            findings.append(
                Finding(
                    filename, para_idx, match.group(), category, suggestion,
                    occurrence=occurrence,
                )
            )
            occurrence += 1

    # Duplicate adjacent lines
    lines = [line.strip() for line in text.splitlines() if _compact(line)]
    for i in range(len(lines) - 1):
        if len(_compact(lines[i])) >= 6 and _compact(lines[i]) == _compact(lines[i + 1]):
            findings.append(
                Finding(
                    filename, para_idx, lines[i], "重复排版",
                    "相邻两行重复出现同一内容。建议删除重复内容。", "high",
                )
            )
            break

    return findings


def detect_word_caption_order(document: Any, filename: str) -> list[Finding]:
    """Detect figure/table caption numbering order by paragraph sequence."""
    findings: list[Finding] = []
    for prefix, category in [("图", "图序/版式"), ("表", "表序/版式")]:
        caps: list[tuple[int, int, str]] = []
        for para_idx, para in enumerate(document.paragraphs, start=1):
            text = para.text.strip()
            match = re.match(rf"^{prefix}\s*([0-9]+)\s+", text)
            if match and len(text) <= 80:
                num = int(match.group(1))
                caps.append((num, para_idx, text))

        previous: tuple[int, int, str] | None = None
        for item in caps:
            if previous and item[0] < previous[0]:
                findings.append(
                    Finding(
                        filename, item[1],
                        f"auto:caption-order:{prefix}:{item[1]}:{item[0]}",
                        category,
                        f"段落顺序检查显示{prefix}{item[0]}编号相对前一处{prefix}{previous[0]}回退。建议核查图表题位置与正文引用顺序。",
                        severity="medium", source="text-rule",
                    )
                )
            previous = item
    return findings


def detect_word_reference_sequence(document: Any, filename: str) -> list[Finding]:
    """Detect reference numbering continuity."""
    refs: list[tuple[int, int, str]] = []
    in_ref_section = False
    for para_idx, para in enumerate(document.paragraphs, start=1):
        for line in para.text.splitlines():
            stripped = line.strip()
            if re.match(r"^参考文献\s*$", stripped):
                in_ref_section = True
                continue
            if not in_ref_section:
                continue
            # Support full-width ［1］, half-width [1], and parentheses (1)
            match = re.match(r"^[\[［(]([0-9]+)[\]］)]", stripped)
            if match:
                refs.append((int(match.group(1)), para_idx, stripped))

    if len(refs) < 2:
        return []

    nums = [item[0] for item in refs]
    duplicates = len(nums) != len(set(nums))
    continuous = nums == list(range(nums[0], nums[-1] + 1))
    if continuous and not duplicates:
        return []

    first = refs[0]
    issues: list[str] = []
    if duplicates:
        issues.append("重复")
    if not continuous:
        issues.append("跳号或不连续")

    return [
        Finding(
            filename, first[1], first[2][:20], "参考文献编号",
            f"参考文献编号可能存在{'、'.join(issues)}。建议按正文引用顺序核对编号连续性。",
            severity="high", source="reference-sequence",
        )
    ]


def detect_word_citation_order(document: Any, filename: str) -> list[Finding]:
    """Detect missing figure/table numbers in first in-text citations."""
    findings: list[Finding] = []
    for prefix, label in [("图", "图"), ("表", "表")]:
        cites: list[tuple[int, int]] = []
        range_covered: set[int] = set()
        for para_idx, para in enumerate(document.paragraphs, start=1):
            text = para.text
            # Detect range citations like "图1—5" and mark covered numbers
            for rmatch in re.finditer(rf"{prefix}\s*([0-9]+)\s*[—~\-～]\s*([0-9]+)", text):
                start_n, end_n = int(rmatch.group(1)), int(rmatch.group(2))
                range_covered.update(range(start_n, end_n + 1))
            for match in re.finditer(rf"{prefix}\s*([0-9]+)", text):
                num = int(match.group(1))
                cites.append((para_idx, num))
        if len(cites) < 2:
            continue
        seen: set[int] = set()
        first_cites: list[tuple[int, int]] = []
        for para_idx, num in cites:
            if num not in seen:
                seen.add(num)
                first_cites.append((para_idx, num))
        if len(first_cites) >= 2:
            nums = [n for _, n in first_cites]
            expected = list(range(min(nums), max(nums) + 1))
            missing = [n for n in expected if n not in nums and n not in range_covered]
            if missing:
                findings.append(
                    Finding(
                        filename, first_cites[0][0],
                        f"auto:{label}-citation-order",
                        f"{label}序/版式",
                        f"正文首次引用的{label}编号不连续，缺少{label}{missing}。建议核查{label}顺序与正文引用是否一致。",
                        severity="medium", source="text-rule",
                    )
                )
    return findings


def detect_placeholders(document: Any, filename: str) -> list[Finding]:
    """Detect template placeholder text that hasn't been filled in."""
    findings: list[Finding] = []
    placeholder_patterns: list[tuple[str, str, str]] = [
        (r"请输入[标题名称作者单位摘要关键词]", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
        (r"此处插入[图表公式]", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
        (r"图[注表]?待[补填]", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
        (r"数据待[补填更新]", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
        (r"\bXXX\b", "内容/占位符", "检测到占位符XXX，请替换为实际内容。"),
        (r"\bTBD\b", "内容/占位符", "检测到占位符TBD，请替换为实际内容。"),
        (r"\bplaceholder\b", "内容/占位符", "检测到占位符placeholder，请替换为实际内容。"),
        (r"Lorem\s+ipsum", "内容/占位符", "检测到Lorem ipsum占位文本，请删除或替换。"),
        (r"（待[填补]）", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
        (r"<待[补填]>", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
        (r"\[请?插入?[一-龥]*\]", "内容/占位符", "检测到占位符文本，请替换为实际内容。"),
    ]
    for para_idx, para in enumerate(document.paragraphs, start=1):
        text = para.text
        if not text.strip():
            continue
        for regex, category, suggestion in placeholder_patterns:
            occurrence = 0
            for match in re.finditer(regex, text, re.IGNORECASE):
                findings.append(
                    Finding(
                        filename, para_idx, match.group(), category, suggestion,
                        severity="high", source="placeholder-rule", occurrence=occurrence,
                    )
                )
                occurrence += 1
    return findings


def detect_superscript_errors(document: Any, filename: str) -> list[Finding]:
    """Detect missing superscripts for common scientific patterns (e.g., R2, m2)."""
    findings: list[Finding] = []
    # Patterns where a digit should be superscript: unit², R², etc.
    # Only include common scientific units and symbols; avoid rare matches.
    superscript_patterns = [
        (r"[Rm]2\b", "单位上标", "2"),
        (r"[Rm]3\b", "单位上标", "3"),
        (r"cm2\b", "单位上标", "2"),
        (r"cm3\b", "单位上标", "3"),
        (r"m2\b", "单位上标", "2"),
        (r"m3\b", "单位上标", "3"),
        (r"km2\b", "单位上标", "2"),
        (r"km3\b", "单位上标", "3"),
        (r"mm2\b", "单位上标", "2"),
        (r"mm3\b", "单位上标", "3"),
        (r"hm2\b", "单位上标", "2"),
        (r"h2\b", "单位上标", "2"),
        (r"s2\b", "单位上标", "2"),
        (r"d2\b", "单位上标", "2"),
        (r"a2\b", "单位上标", "2"),
    ]
    # Chemical subscript patterns
    chemical_patterns = [
        (r"H2O\b", "化学式", "H₂O"),
        (r"CO2\b", "化学式", "CO₂"),
        (r"SO2\b", "化学式", "SO₂"),
        (r"NO2\b", "化学式", "NO₂"),
        (r"CH4\b", "化学式", "CH₄"),
        (r"NH3\b", "化学式", "NH₃"),
    ]

    for para_idx, para in enumerate(document.paragraphs, start=1):
        full_text, runs = _build_para_text_with_runs(para)
        if not full_text.strip():
            continue

        for regex, desc, expected in superscript_patterns:
            for match in re.finditer(regex, full_text, re.IGNORECASE):
                # Check if the digit is superscripted
                digit_offset = match.end() - 1
                run = _find_run_at_offset(runs, digit_offset)
                if run is None:
                    continue
                rPr = run._r.get_or_add_rPr()
                vert = rPr.find(qn("w:vertAlign"))
                is_sup = vert is not None and vert.get(qn("w:val")) == "superscript"
                if not is_sup:
                    findings.append(
                        Finding(
                            filename, para_idx, match.group(),
                            "公式/上标格式",
                            f"格式检查显示，{match.group()} 中的数字未显示为上标。建议改为{expected}上标格式。",
                            severity="medium", source="word-format",
                        )
                    )

        for regex, desc, expected in chemical_patterns:
            for match in re.finditer(regex, full_text, re.IGNORECASE):
                # Check if the digit is subscripted
                digit_offset = match.start() + 1
                run = _find_run_at_offset(runs, digit_offset)
                if run is None:
                    continue
                rPr = run._r.get_or_add_rPr()
                vert = rPr.find(qn("w:vertAlign"))
                is_sub = vert is not None and vert.get(qn("w:val")) == "subscript"
                if not is_sub:
                    findings.append(
                        Finding(
                            filename, para_idx, match.group(),
                            "公式/下标格式",
                            f"格式检查显示，{match.group()} 中的数字未显示为下标。建议改为{expected}格式。",
                            severity="medium", source="word-format",
                        )
                    )
    return findings


def detect_common_typos(document: Any, filename: str) -> list[Finding]:
    """Detect common Chinese and English typos beyond the hardcoded list."""
    findings: list[Finding] = []

    # Chinese typo patterns (heuristic regexes)
    # NOTE: 的/地/得 rules are deliberately omitted because the adverbial
    # particle 地 before a verb is grammatically correct (e.g. 快速地完成).
    chinese_typos: list[tuple[str, str, str]] = [
        (r"走的快", "文字/用词", "“走的快”疑为“走得快”，补语前应用“得”。"),
        (r"跑的快", "文字/用词", "“跑的快”疑为“跑得快”，补语前应用“得”。"),
        (r"在次", "文字/用词", "“在次”疑为“再次”。"),
        (r"做用", "文字/用词", "“做用”疑为“作用”。"),
        (r"做为", "文字/用词", "“做为”疑为“作为”。"),
        (r"好象", "文字/用词", "“好象”应为“好像”。"),
        (r"图象", "文字/用词", "“图象”在现代汉语中通常应为“图像”。"),
        (r"帐号", "文字/用词", "“帐号”应为“账号”。"),
        (r"帐本", "文字/用词", "“帐本”应为“账本”。"),
    ]

    # English typo patterns
    english_typos: list[tuple[str, str, str]] = [
        (r"\bteh\b", "文字/拼写", "英文拼写错误：teh → the。"),
        (r"\badn\b", "文字/拼写", "英文拼写错误：adn → and。"),
        (r"\boccurence\b", "文字/拼写", "英文拼写错误：occurence → occurrence。"),
        (r"\bseperate\b", "文字/拼写", "英文拼写错误：seperate → separate。"),
        (r"\bdefinately\b", "文字/拼写", "英文拼写错误：definately → definitely。"),
        (r"\brecieve\b", "文字/拼写", "英文拼写错误：recieve → receive。"),
        (r"\boccured\b", "文字/拼写", "英文拼写错误：occured → occurred。"),
        (r"\bgoverment\b", "文字/拼写", "英文拼写错误：goverment → government。"),
        (r"\benviroment\b", "文字/拼写", "英文拼写错误：enviroment → environment。"),
        (r"\bwich\b", "文字/拼写", "英文拼写错误：wich → which。"),
        (r"\buntill\b", "文字/拼写", "英文拼写错误：untill → until。"),
    ]

    all_typos = chinese_typos + english_typos

    for para_idx, para in enumerate(document.paragraphs, start=1):
        text = para.text
        if not text.strip():
            continue
        for regex, category, suggestion in all_typos:
            occurrence = 0
            for match in re.finditer(regex, text):
                findings.append(
                    Finding(
                        filename, para_idx, match.group(), category, suggestion,
                        occurrence=occurrence, source="typo-rule",
                    )
                )
                occurrence += 1
    return findings


def detect_inconsistent_compounds(document: Any, filename: str) -> list[Finding]:
    """Detect if the same compound word is written in inconsistent forms across the document.

    NOTE: Disabled from auto-finding generation because noun-phrase and adjectival forms
    (e.g., "land use" vs "land-use") are grammatically correct differences, not errors.
    The raw hints are still available for AI editorial review but are not emitted as
    automatic annotations to avoid hallucination.
    """
    # Collect full document text
    full_text = "\n".join(para.text for para in document.paragraphs)
    if not full_text.strip():
        return []

    # Build hint list for review_context.md consumption only (no automatic findings)
    compound_groups: list[list[str]] = [
        ["co-operation", "cooperation"],
        ["e-mail", "email", "E-mail"],
        ["on-line", "online"],
        ["work flow", "workflow"],
        ["data base", "database"],
        ["land use", "land-use", "landuse"],
        ["soil erosion", "soil-erosion"],
        ["high performance", "high-performance"],
        ["water quality", "water-quality"],
        ["climate change", "climate-change"],
        ["carbon sequestration", "carbon-sequestration"],
        ["soil moisture", "soil-moisture"],
        ["run off", "runoff", "run-off"],
        ["check list", "checklist"],
        ["feed back", "feedback"],
        ["set up", "setup"],
        ["out put", "output"],
        ["in put", "input"],
    ]

    hints: list[str] = []
    for group in compound_groups:
        found_variants = [variant for variant in group if variant.lower() in full_text.lower()]
        if len(found_variants) >= 2:
            hints.append(f"{', '.join(found_variants)}")

    # We intentionally do NOT return Finding objects here to avoid false-positive annotations.
    # If needed in the future, write hints to a sidecar file for the AI reviewer.
    return []


def dedupe_findings(findings: list[Finding]) -> list[Finding]:
    seen: set[tuple[object, ...]] = set()
    unique: list[Finding] = []
    for item in findings:
        key = (item.file, item.page, item.target, item.category, item.occurrence, item.fallback_rect, item.source)
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)
    return unique
