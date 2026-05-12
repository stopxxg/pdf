"""Word (.docx) → Markdown converter with format markup.

Extracts text from .docx preserving italic, bold, superscript, and subscript
information from OOXML run properties, and generates structured Markdown with
[[PARA=N]] block markers for AI review.

Also extracts tables and headers/footers for comprehensive review.
"""

from __future__ import annotations

from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _run_has_property(run: Any, prop_name: str, prop_val: str | None = None) -> bool:
    """Check if a run has a specific XML property."""
    rPr = run._r.get_or_add_rPr()
    elem = rPr.find(qn(f"w:{prop_name}"))
    if elem is None:
        return False
    if prop_val is None:
        return True
    return elem.get(qn("w:val")) == prop_val


def _is_italic(run: Any) -> bool:
    if run.font.italic:
        return True
    return _run_has_property(run, "i")


def _is_bold(run: Any) -> bool:
    if run.font.bold:
        return True
    return _run_has_property(run, "b")


def _is_superscript(run: Any) -> bool:
    return _run_has_property(run, "vertAlign", "superscript")


def _is_subscript(run: Any) -> bool:
    return _run_has_property(run, "vertAlign", "subscript")


def _format_run_text(run: Any) -> str:
    text = run.text
    if not text:
        return ""
    if _is_superscript(run):
        text = f"<sup>{text}</sup>"
    elif _is_subscript(run):
        text = f"<sub>{text}</sub>"
    if _is_italic(run):
        text = f"<i>{text}</i>"
    if _is_bold(run):
        text = f"<b>{text}</b>"
    return text


def _paragraph_to_text(paragraph: Any) -> str:
    parts = [_format_run_text(run) for run in paragraph.runs if run.text]
    return "".join(parts)


def _cell_to_text(cell: Any) -> str:
    """Extract text from a table cell, preserving markup."""
    parts: list[str] = []
    for para in cell.paragraphs:
        text = _paragraph_to_text(para)
        if text.strip():
            parts.append(text)
    return " ".join(parts) if parts else ""


def _table_to_markdown(table: Any, table_index: int, document: Any) -> str:
    """Convert a single table to Markdown pipe table format."""
    lines: list[str] = []
    lines.append(f"\n[[TABLE={table_index}]]\n")

    # Try to find caption: paragraph immediately before the table
    caption = ""
    try:
        # Locate the table element in the document body
        tbl_elem = table._tbl
        body = document.element.body
        children = list(body)
        tbl_idx = children.index(tbl_elem)
        if tbl_idx > 0:
            prev = children[tbl_idx - 1]
            if prev.tag.endswith("p"):
                prev_text = "".join(t.text or "" for t in prev.iter())
                if any(kw in prev_text for kw in ("表", "Table", "图", "Figure")):
                    caption = prev_text.strip()
    except Exception:
        pass

    if caption:
        lines.append(f"**Caption**: {caption}\n")

    rows = table.rows
    if not rows:
        lines.append("(empty table)\n")
        return "\n".join(lines)

    # Header separator
    header_sep_drawn = False

    for r_idx, row in enumerate(rows):
        cells = row.cells
        cell_texts = [_cell_to_text(cell) for cell in cells]
        # Escape pipe characters in cell text
        cell_texts = [t.replace("|", "\\|") for t in cell_texts]
        lines.append("| " + " | ".join(cell_texts) + " |")

        if r_idx == 0:
            lines.append("|" + "|".join(" --- " for _ in cells) + "|")

    lines.append("")
    return "\n".join(lines)


def _tables_to_markdown(document: Any) -> str:
    """Convert all tables in the document to Markdown."""
    if not document.tables:
        return ""
    parts: list[str] = ["\n# Tables\n"]
    for idx, table in enumerate(document.tables, start=1):
        parts.append(_table_to_markdown(table, idx, document))
    return "\n".join(parts)


def _headers_footers_to_text(document: Any) -> str:
    """Extract text from all sections' headers and footers."""
    parts: list[str] = []
    for sec_idx, section in enumerate(document.sections, start=1):
        sec_parts: list[str] = []

        # Header
        header = section.header
        if header is not None:
            header_texts = [_paragraph_to_text(p) for p in header.paragraphs if _paragraph_to_text(p).strip()]
            if header_texts:
                sec_parts.append("[Header]\n" + "\n".join(header_texts))

        # Footer
        footer = section.footer
        if footer is not None:
            footer_texts = [_paragraph_to_text(p) for p in footer.paragraphs if _paragraph_to_text(p).strip()]
            if footer_texts:
                sec_parts.append("[Footer]\n" + "\n".join(footer_texts))

        if sec_parts:
            parts.append(f"## Section {sec_idx}\n" + "\n".join(sec_parts))

    if not parts:
        return ""
    return "\n# Headers and Footers\n\n" + "\n\n".join(parts) + "\n"


def docx_to_markdown(docx_path: Path, block_size: int = 30) -> str:
    """Convert an entire .docx to Markdown with block markers and format tags."""
    document = Document(str(docx_path))
    parts: list[str] = []
    para_idx = 0

    for para in document.paragraphs:
        text = _paragraph_to_text(para)
        if not text.strip():
            continue
        para_idx += 1
        if (para_idx - 1) % block_size == 0:
            block_no = (para_idx - 1) // block_size + 1
            parts.append(f"\n[[PARA={block_no}]]\n")
        parts.append(text)

    # Append tables at the end
    tables_md = _tables_to_markdown(document)
    if tables_md:
        parts.append(tables_md)

    return "\n\n".join(parts)


def write_text_artifacts(docx_path: Path, artifact_dir: Path, block_size: int = 30) -> Path:
    """Write per-paragraph, per-block, table, header/footer text artifacts for AI review."""
    paper_dir = artifact_dir / docx_path.stem
    paper_dir.mkdir(parents=True, exist_ok=True)

    document = Document(str(docx_path))
    chunks: list[str] = []
    block_texts: list[tuple[int, str]] = []
    para_idx = 0

    for para in document.paragraphs:
        text = _paragraph_to_text(para)
        if not text.strip():
            continue
        para_idx += 1
        chunks.append(f"===== PARA {para_idx} =====\n{text}\n")
        block_no = (para_idx - 1) // block_size + 1
        block_texts.append((block_no, text))

    (paper_dir / "fulltext.txt").write_text("\n".join(chunks), encoding="utf-8")

    # Group by block
    blocks: defaultdict[int, list[str]] = defaultdict(list)
    for block_no, text in block_texts:
        blocks[block_no].append(text)

    page_dir = paper_dir / "page_paras"
    page_dir.mkdir(exist_ok=True)
    for block_no, texts in sorted(blocks.items()):
        content = f"===== BLOCK {block_no} =====\n" + "\n\n".join(texts)
        (page_dir / f"page_{block_no:03d}.txt").write_text(content, encoding="utf-8")

    md = docx_to_markdown(docx_path, block_size=block_size)
    (paper_dir / "fulltext.md").write_text(md, encoding="utf-8")

    # Write tables
    tables_md = _tables_to_markdown(document)
    if tables_md:
        (paper_dir / "tables.md").write_text(tables_md, encoding="utf-8")
    else:
        (paper_dir / "tables.md").write_text("# Tables\n\nNo tables found in this document.\n", encoding="utf-8")

    # Write headers/footers
    hf_text = _headers_footers_to_text(document)
    if hf_text:
        (paper_dir / "headers_footers.txt").write_text(hf_text, encoding="utf-8")
    else:
        (paper_dir / "headers_footers.txt").write_text("# Headers and Footers\n\nNone found.\n", encoding="utf-8")

    return paper_dir
