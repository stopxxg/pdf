#!/usr/bin/env python3
"""Low-output Word (.docx) proofreading pipeline for Chinese academic journals.

Mirrors the PDF pipeline structure but works with .docx files, adding native
Word comment annotations instead of red rectangles.
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ModuleNotFoundError:
    print("Error: python-docx is required. Install: pip install python-docx>=0.8.11", file=sys.stderr)
    sys.exit(1)

from word_module.word_text_converter import write_text_artifacts
from word_module.word_annotator import annotate_docx as _annotate_docx_impl
from word_module.word_rule_detectors import (
    detect_word_stat_symbol_style,
    detect_word_script_style,
    detect_word_text_rules,
    detect_word_caption_order,
    detect_word_reference_sequence,
    detect_word_citation_order,
    dedupe_findings,
    Finding,
)


@dataclass(frozen=True)
class WordFinding:
    # Same fields as PDF Finding for cross-compatibility
    file: str
    page: int
    target: str
    category: str
    suggestion: str
    severity: str = "medium"
    source: str = "rule"
    occurrence: int = 0
    fallback_rect: tuple[float, float, float, float] | None = None
    end_of_doc: bool = False


def docxs_under(root: Path, output_dir: Path) -> list[Path]:
    output_dir = output_dir.resolve()
    docxs: list[Path] = []
    for path in sorted(root.rglob("*")):
        if path.suffix.lower() != ".docx" or not path.is_file():
            continue
        try:
            if output_dir in path.resolve().parents:
                continue
        except FileNotFoundError:
            pass
        docxs.append(path)
    return docxs


def collect_candidates(docx: Path, output_dir: Path) -> dict[str, object]:
    artifact_dir = output_dir / "_artifacts"
    candidate_dir = output_dir / "_candidates"
    candidate_dir.mkdir(parents=True, exist_ok=True)
    result: dict[str, object] = {
        "file": str(docx),
        "status": "skipped",
        "reason": "",
        "pages": 0,
        "candidate_count": 0,
        "candidate_json": "",
        "artifact_dir": "",
    }
    try:
        document = Document(str(docx))
    except Exception as exc:
        result["reason"] = f"open failed: {exc}"
        return result

    try:
        para_count = len(document.paragraphs)
        result["pages"] = para_count
        write_text_artifacts(docx, artifact_dir)

        findings: list[Finding] = []
        findings.extend(detect_word_stat_symbol_style(document, docx.name))
        findings.extend(detect_word_script_style(document, docx.name))
        findings.extend(detect_word_caption_order(document, docx.name))
        findings.extend(detect_word_reference_sequence(document, docx.name))
        findings.extend(detect_word_citation_order(document, docx.name))

        for para_idx, para in enumerate(document.paragraphs, start=1):
            text = para.text
            if text.strip():
                findings.extend(detect_word_text_rules(docx.name, para_idx, text))

        findings = dedupe_findings(findings)
        candidate_path = candidate_dir / f"{docx.stem}.findings.json"
        candidate_path.write_text(
            json.dumps([asdict(item) for item in findings], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        result.update({
            "status": "done",
            "candidate_count": len(findings),
            "candidate_json": str(candidate_path),
            "artifact_dir": str(artifact_dir / docx.stem),
        })
        return result
    except Exception as exc:
        result["reason"] = f"scan failed: {exc}"
        return result


def load_findings(path: Path) -> list[WordFinding]:
    data = json.loads(path.read_text(encoding="utf-8"))
    findings: list[WordFinding] = []
    for item in data:
        rect = item.get("fallback_rect")
        findings.append(
            WordFinding(
                file=item["file"],
                page=int(item["page"]),
                target=item["target"],
                category=item["category"],
                suggestion=item["suggestion"],
                severity=item.get("severity", "medium"),
                source=item.get("source", "reviewed"),
                occurrence=int(item.get("occurrence", 0)),
                fallback_rect=tuple(rect) if rect else None,
                end_of_doc=bool(item.get("end_of_doc", False)),
            )
        )
    return findings


def annotate_docx(docx: Path, output_dir: Path, findings: list[WordFinding], same_name: bool) -> dict[str, object]:
    result = _annotate_docx_impl(docx, findings, output_dir / docx.name, same_name=same_name)
    # Convert dataclass objects in missed list to plain dicts for JSON serialization
    if result.get("missed"):
        result["missed"] = [
            {**item, "finding": asdict(item["finding"])} if hasattr(item.get("finding"), "__dataclass_fields__") else item
            for item in result["missed"]
        ]
    return result


def append_log(output_dir: Path, entries: list[dict[str, object]]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / "batch_proofread_log.json"
    prior: list[dict[str, object]] = []
    if path.exists():
        try:
            loaded = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(loaded, list):
                prior = loaded
        except Exception:
            prior = []
    replaced = {Path(str(item.get("file", ""))).name for item in entries}
    prior = [item for item in prior if Path(str(item.get("file", ""))).name not in replaced]
    prior.extend(entries)
    path.write_text(json.dumps(prior, ensure_ascii=False, indent=2), encoding="utf-8")


def resolve_output(root: Path, output: Path | None) -> Path:
    return output.resolve() if output else (root / "annotated_docxs").resolve()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, required=True, help="Folder containing source .docx files.")
    parser.add_argument("--output", type=Path, help="Output folder. Defaults to ROOT/annotated_docxs.")
    parser.add_argument("--mode", choices=["scan", "annotate"], default="scan")
    parser.add_argument("--findings-json", type=Path, help="Reviewed findings JSON for annotate mode.")
    parser.add_argument("--same-name", action="store_true", help="Keep output filenames unchanged.")
    parser.add_argument("--limit", type=int, default=0, help="Maximum .docx files to process; 0 means all.")
    parser.add_argument("--verbose", action="store_true", help="Print full JSON results instead of compact summary.")
    args = parser.parse_args()

    root = args.root.resolve()
    output_dir = resolve_output(root, args.output)
    docxs = docxs_under(root, output_dir)
    if args.limit > 0:
        docxs = docxs[: args.limit]

    results: list[dict[str, object]] = []
    if args.mode == "scan":
        for docx in docxs:
            item = collect_candidates(docx, output_dir)
            results.append(item)
            print(json.dumps({"phase": "scan", "file": docx.name, "status": item["status"], "candidates": item.get("candidate_count", 0)}, ensure_ascii=False))

    if args.mode == "annotate":
        if args.findings_json:
            findings = load_findings(args.findings_json)
        else:
            raise SystemExit("--findings-json is required in annotate mode")

        annotate_results = []
        for docx in docxs:
            item = annotate_docx(docx, output_dir, findings, same_name=args.same_name)
            annotate_results.append(item)
            print(
                json.dumps(
                    {
                        "phase": "annotate",
                        "file": docx.name,
                        "status": item["status"],
                        "annotations": item.get("findings_annotated", 0),
                        "missed": len(item.get("missed", [])),
                    },
                    ensure_ascii=False,
                )
            )
        append_log(output_dir, annotate_results)
        results = annotate_results

    if args.verbose:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        print(
            json.dumps(
                {
                    "root": str(root),
                    "output": str(output_dir),
                    "mode": args.mode,
                    "docxs": len(docxs),
                    "done": sum(1 for item in results if str(item.get("status", "")).startswith("done")),
                    "skipped": sum(1 for item in results if item.get("status") == "skipped"),
                    "total_annotations": sum(int(item.get("findings_annotated", 0) or 0) for item in results),
                },
                ensure_ascii=False,
            )
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
