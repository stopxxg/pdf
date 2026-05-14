#!/usr/bin/env python3
"""High-precision Word (.docx) proofreading pipeline for Chinese academic journals.

Extracts rich text with OOXML markup, runs rule checks, generates review_context.md
for deep AI editorial review, and inserts native Word comments for annotations.
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ModuleNotFoundError:
    print("Error: python-docx is required. Install: pip install python-docx>=0.8.11", file=sys.stderr)
    sys.exit(1)

from word_module.word_text_converter import write_text_artifacts
from word_module.word_annotator import annotate_docx as _annotate_docx_impl
from word_module.word_rule_detectors import (
    detect_word_stat_symbol_style,
    detect_word_script_style,
    detect_word_text_rules,
    detect_word_caption_order,
    detect_word_reference_sequence,
    detect_word_citation_order,
    detect_placeholders,
    detect_superscript_errors,
    detect_common_typos,
    detect_inconsistent_compounds,
    dedupe_findings,
    Finding,
)


@dataclass(frozen=True)
class WordFinding:
    file: str
    page: int
    target: str
    category: str
    suggestion: str
    severity: str = "medium"
    source: str = "rule"
    occurrence: int = 0
    fallback_rect: tuple[float, float, float, float] | None = None
    end_of_doc: bool = False


def docxs_under(root: Path, output_dir: Path) -> list[Path]:
    output_dir = output_dir.resolve()
    docxs: list[Path] = []
    for path in sorted(root.rglob("*")):
        if path.suffix.lower() != ".docx" or not path.is_file():
            continue
        try:
            if output_dir in path.resolve().parents:
                continue
        except FileNotFoundError:
            pass
        docxs.append(path)
    return docxs


def prepare_document(docx: Path, output_dir: Path) -> dict[str, object]:
    """Extract text, run rule checks, and generate review context for AI editorial review."""
    artifact_dir = output_dir / "_artifacts"
    candidate_dir = output_dir / "_candidates"
    candidate_dir.mkdir(parents=True, exist_ok=True)

    result: dict[str, object] = {
        "file": str(docx),
        "status": "skipped",
        "reason": "",
        "paragraphs": 0,
        "candidate_count": 0,
        "candidate_json": "",
        "artifact_dir": "",
    }

    try:
        document = Document(str(docx))
    except Exception as exc:
        result["reason"] = f"open failed: {exc}"
        return result

    try:
        para_count = len(document.paragraphs)
        result["paragraphs"] = para_count
        write_text_artifacts(docx, artifact_dir)

        # Run all rule-based detectors
        findings: list[Finding] = []
        findings.extend(detect_word_stat_symbol_style(document, docx.name))
        findings.extend(detect_word_script_style(document, docx.name))
        findings.extend(detect_word_caption_order(document, docx.name))
        findings.extend(detect_word_reference_sequence(document, docx.name))
        findings.extend(detect_word_citation_order(document, docx.name))
        findings.extend(detect_placeholders(document, docx.name))
        findings.extend(detect_superscript_errors(document, docx.name))
        findings.extend(detect_common_typos(document, docx.name))
        findings.extend(detect_inconsistent_compounds(document, docx.name))

        for para_idx, para in enumerate(document.paragraphs, start=1):
            text = para.text
            if text.strip():
                findings.extend(detect_word_text_rules(docx.name, para_idx, text))

        findings = dedupe_findings(findings)

        # Write initial findings JSON
        candidate_path = candidate_dir / f"{docx.stem}.findings.json"
        candidate_path.write_text(
            json.dumps([asdict(item) for item in findings], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        # Build review_context.md
        review_context = _build_review_context(docx, document, artifact_dir, findings)
        review_path = artifact_dir / docx.stem / "review_context.md"
        review_path.write_text(review_context, encoding="utf-8")

        result.update({
            "status": "done",
            "candidate_count": len(findings),
            "candidate_json": str(candidate_path),
            "artifact_dir": str(artifact_dir / docx.stem),
            "review_context": str(review_path),
        })
        return result
    except Exception as exc:
        result["reason"] = f"prepare failed: {exc}"
        return result


def _build_review_context(docx: Path, document: Any, artifact_dir: Path, findings: list[Finding]) -> str:
    """Build a unified review context document for AI editorial review."""
    parts: list[str] = []

    parts.append(f"# Review Context: {docx.name}")
    parts.append(f"- Paragraphs: {len(document.paragraphs)}")
    parts.append("")

    # Read fulltext.md
    fulltext_md_path = artifact_dir / docx.stem / "fulltext.md"
    if fulltext_md_path.exists():
        fulltext = fulltext_md_path.read_text(encoding="utf-8")
        parts.append("# Full Text (with markup)\n")
        parts.append(fulltext)
        parts.append("")

    # Tables content
    tables_md_path = artifact_dir / docx.stem / "tables.md"
    if tables_md_path.exists():
        tables_text = tables_md_path.read_text(encoding="utf-8")
        parts.append("# Tables\n")
        parts.append(tables_text)
        parts.append("")

    # Headers and footers
    hf_path = artifact_dir / docx.stem / "headers_footers.txt"
    if hf_path.exists():
        hf_text = hf_path.read_text(encoding="utf-8")
        parts.append("# Headers and Footers\n")
        parts.append(hf_text)
        parts.append("")

    # Auto-detected rule findings
    parts.append("# Auto-Detected Rule Findings\n")
    if findings:
        parts.append(f"Total: {len(findings)} findings\n")

        # Group by category for structured display
        by_category: dict[str, list[Finding]] = {}
        for f in findings:
            by_category.setdefault(f.category, []).append(f)

        # High-priority categories first
        priority_order = [
            "内容/占位符",
            "公式/统计符号正斜体",
            "公式/上标格式",
            "公式/下标格式",
            "文字/拼写",
            "文字/用词",
            "文字/标点",
            "术语/统一",
            "参考文献编号",
            "参考文献格式",
            "图序/版式",
            "表序/版式",
        ]
        sorted_categories = sorted(
            by_category.keys(),
            key=lambda c: (priority_order.index(c) if c in priority_order else 999, c),
        )

        for category in sorted_categories:
            items = by_category[category]
            parts.append(f"## {category} ({len(items)})\n")
            for item in items:
                parts.append(f"- **Paragraph {item.page}**: {item.suggestion}")
                if item.target and not item.target.startswith("auto:"):
                    parts.append(f"  - Target: `{item.target}`")
            parts.append("")
    else:
        parts.append("No rule-based findings detected.\n")

    # Paragraphs needing attention (figures, tables, formulas)
    attention_paras: set[int] = set()
    for para_idx, para in enumerate(document.paragraphs, start=1):
        text = para.text
        if any(kw in text for kw in ("图", "表", "公式", "Fig.", "Table", "Equation", "式中")):
            attention_paras.add(para_idx)
        for f in findings:
            if f.page == para_idx:
                attention_paras.add(para_idx)

    parts.append("# Paragraphs Requiring Close Attention\n")
    if attention_paras:
        for p in sorted(attention_paras)[:50]:  # Limit to avoid overwhelming
            text = document.paragraphs[p - 1].text[:120].replace("\n", " ")
            parts.append(f"- Paragraph {p}: `{text}{'...' if len(document.paragraphs[p - 1].text) > 120 else ''}`")
        if len(attention_paras) > 50:
            parts.append(f"- ... and {len(attention_paras) - 50} more paragraphs")
        parts.append("")
    else:
        parts.append("No specific paragraphs flagged.\n")

    parts.append("# Review Instructions")
    parts.append("1. Read the full text above sentence by sentence.")
    parts.append("2. Check front matter: title, authors, abstract structure, keywords, CLC, DOI, funding.")
    parts.append("3. Check text and logic: repeated phrasing, dangling referents, conclusion/data mismatch, term consistency.")
    parts.append("4. Check spelling and grammar:")
    parts.append("   - Chinese homophones: 的/地/得, 在/再, 做/作, 象/像, 帐号/账号, 其它/其他.")
    parts.append("   - English typos and OCR artifacts.")
    parts.append("   - Grammar: subject-verb disagreement, missing articles, incorrect measure words, redundant words.")
    parts.append("5. Check tense consistency: results sections should use past tense; general facts may use present tense.")
    parts.append("6. Check template placeholders: scan for unfilled text such as 请输入标题, XXX, 待补充, TBD, placeholder, 图注待补.")
    parts.append("7. Check punctuation: no full-width colons in URLs, correct comma/period usage, consistent parentheses.")
    parts.append("8. Check numerals and units: consistent spacing, correct range connectors, unbroken statistical expressions.")
    parts.append("9. Check formulas and symbols: italic variables, upright operators/units, correct subscripts/superscripts.")
    parts.append("   - Use `<i>` markup in the text to verify italics; use `<sub>` / `<sup>` to verify scripts.")
    parts.append("10. Check tables: empty cells, inconsistent decimal places, incorrect units, data-text consistency.")
    parts.append("11. Check headers/footers: consistent page numbers, wrong journal names, mixed full-width/half-width punctuation.")
    parts.append("12. Check references: GB/T 7714 format, volume/issue completeness, citation-reference consistency, bilingual refs.")
    parts.append("13. Check hyphenation and compound-word consistency: land use vs land-use vs landuse, email vs e-mail, etc.")
    parts.append("14. Check capitalization of technical terms: GIS vs Gis, NDVI vs Ndvi, RUSLE vs Rusle.")
    parts.append("15. Check abbreviations are defined on first use.")
    parts.append("16. Cross-check auto-detected findings above; confirm or dismiss each one.")
    parts.append("17. For every confirmed issue, produce a finding with: paragraph number, target text, category, severity, suggestion.")

    return "\n".join(parts) + "\n"


def load_findings(path: Path) -> list[WordFinding]:
    data = json.loads(path.read_text(encoding="utf-8"))
    findings: list[WordFinding] = []
    for item in data:
        rect = item.get("fallback_rect")
        findings.append(
            WordFinding(
                file=item["file"],
                page=int(item["page"]),
                target=item["target"],
                category=item["category"],
                suggestion=item["suggestion"],
                severity=item.get("severity", "medium"),
                source=item.get("source", "reviewed"),
                occurrence=int(item.get("occurrence", 0)),
                fallback_rect=tuple(rect) if rect else None,
                end_of_doc=bool(item.get("end_of_doc", False)),
            )
        )
    return findings


def annotate_docx(docx: Path, output_dir: Path, findings: list[WordFinding], same_name: bool) -> dict[str, object]:
    result = _annotate_docx_impl(docx, findings, output_dir / docx.name, same_name=same_name)
    if result.get("missed"):
        result["missed"] = [
            {**item, "finding": asdict(item["finding"])} if hasattr(item.get("finding"), "__dataclass_fields__") else item
            for item in result["missed"]
        ]
    return result


def append_log(output_dir: Path, entries: list[dict[str, object]]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / "batch_proofread_log.json"
    prior: list[dict[str, object]] = []
    if path.exists():
        try:
            loaded = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(loaded, list):
                prior = loaded
        except Exception:
            prior = []
    replaced = {Path(str(item.get("file", ""))).name for item in entries}
    prior = [item for item in prior if Path(str(item.get("file", ""))).name not in replaced]
    prior.extend(entries)
    path.write_text(json.dumps(prior, ensure_ascii=False, indent=2), encoding="utf-8")


def resolve_output(root: Path, output: Path | None) -> Path:
    return output.resolve() if output else (root / "annotated_docxs").resolve()


def main() -> int:
    parser = argparse.ArgumentParser(
        description="High-precision Word proofreading pipeline for Chinese academic journals.",
    )
    parser.add_argument("--root", type=Path, help="Folder containing source .docx files. Optional when --file is given.")
    parser.add_argument("--file", type=Path, help="Process a single specific .docx file instead of scanning a folder.")
    parser.add_argument("--output", type=Path, help="Output folder. Defaults to ROOT/annotated_docxs.")
    parser.add_argument("--mode", choices=["prepare", "annotate"], default="prepare",
                        help="prepare: extract text, run rule checks, build review context. "
                             "annotate: apply reviewed findings as native Word comments.")
    parser.add_argument("--findings-json", type=Path, help="Reviewed findings JSON for annotate mode.")
    parser.add_argument("--same-name", action="store_true", help="Keep output filenames unchanged.")
    parser.add_argument("--limit", type=int, default=0, help="Maximum .docx files to process; 0 means all.")
    parser.add_argument("--verbose", action="store_true", help="Print full JSON results instead of compact summary.")
    args = parser.parse_args()

    if args.file:
        docxs = [args.file.resolve()]
        root = args.file.resolve().parent
        output_dir = resolve_output(root, args.output)
    else:
        if not args.root:
            raise SystemExit("Either --root or --file is required.")
        root = args.root.resolve()
        output_dir = resolve_output(root, args.output)
        docxs = docxs_under(root, output_dir)
        if args.limit > 0:
            docxs = docxs[: args.limit]

    results: list[dict[str, object]] = []
    if args.mode == "prepare":
        for docx in docxs:
            item = prepare_document(docx, output_dir)
            results.append(item)
            print(json.dumps({
                "phase": "prepare",
                "file": docx.name,
                "status": item["status"],
                "paragraphs": item.get("paragraphs", 0),
                "candidates": item.get("candidate_count", 0),
                "review_context": item.get("review_context", ""),
            }, ensure_ascii=False))

    if args.mode == "annotate":
        if args.findings_json:
            findings = load_findings(args.findings_json)
        else:
            raise SystemExit("--findings-json is required in annotate mode")

        annotate_results = []
        for docx in docxs:
            item = annotate_docx(docx, output_dir, findings, same_name=args.same_name)
            annotate_results.append(item)
            print(json.dumps({
                "phase": "annotate",
                "file": docx.name,
                "status": item["status"],
                "annotations": item.get("findings_annotated", 0),
                "missed": len(item.get("missed", [])),
            }, ensure_ascii=False))
        append_log(output_dir, annotate_results)
        results = annotate_results

    if args.verbose:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        print(json.dumps({
            "root": str(root),
            "output": str(output_dir),
            "mode": args.mode,
            "docxs": len(docxs),
            "done": sum(1 for item in results if str(item.get("status", "")).startswith("done")),
            "skipped": sum(1 for item in results if item.get("status") == "skipped"),
            "total_annotations": sum(int(item.get("findings_annotated", 0) or 0) for item in results),
        }, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
