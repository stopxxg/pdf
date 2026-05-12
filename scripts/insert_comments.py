#!/usr/bin/env python3
"""Insert review comments into a Word document."""

import sys
import os
# Remove script directory from path to avoid shadowing installed python-docx
script_dir = os.path.dirname(os.path.abspath(__file__))
if script_dir in sys.path:
    sys.path.remove(script_dir)
if '' in sys.path:
    sys.path.remove('')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from datetime import datetime

DOC_PATH = "/Users/iswcxxg/Desktop/ccc/张宇.docx"
AUTHOR = "Editor"
DATE_STR = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

comments_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml">
</w:comments>'''
comments_element = etree.fromstring(comments_xml.encode('utf-8'))

comment_id_counter = [0]

def add_comment_to_xml(comments_element, comment_id, author, date_str, text):
    comment = etree.SubElement(comments_element, qn('w:comment'))
    comment.set(qn('w:id'), str(comment_id))
    comment.set(qn('w:author'), author)
    comment.set(qn('w:date'), date_str)
    p = etree.SubElement(comment, qn('w:p'))
    r = etree.SubElement(p, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')

def add_comment_markers(paragraph, search_text, comment_id):
    """Add commentRangeStart, commentRangeEnd, and commentReference to a paragraph."""
    found = False
    for run in paragraph.runs:
        if search_text.lower() in (run.text or "").lower():
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), str(comment_id))
            run.element.addprevious(comment_start)

            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), str(comment_id))
            run.element.addnext(comment_end)

            ref_run = OxmlElement('w:r')
            ref_rpr = OxmlElement('w:rPr')
            ref_style = OxmlElement('w:rStyle')
            ref_style.set(qn('w:val'), 'CommentReference')
            ref_rpr.append(ref_style)
            ref_run.append(ref_rpr)
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), str(comment_id))
            ref_run.append(comment_ref)
            comment_end.addnext(ref_run)
            found = True
            break

    if not found:
        p_element = paragraph._element
        comment_start = OxmlElement('w:commentRangeStart')
        comment_start.set(qn('w:id'), str(comment_id))
        p_element.insert(0, comment_start)

        comment_end = OxmlElement('w:commentRangeEnd')
        comment_end.set(qn('w:id'), str(comment_id))
        p_element.append(comment_end)

        ref_run = OxmlElement('w:r')
        ref_rpr = OxmlElement('w:rPr')
        ref_style = OxmlElement('w:rStyle')
        ref_style.set(qn('w:val'), 'CommentReference')
        ref_rpr.append(ref_style)
        ref_run.append(ref_rpr)
        comment_ref = OxmlElement('w:commentReference')
        comment_ref.set(qn('w:id'), str(comment_id))
        ref_run.append(comment_ref)
        p_element.append(ref_run)

def attach_comments_part(doc, comments_element):
    comments_bytes = etree.tostring(comments_element, xml_declaration=True, encoding='UTF-8', standalone=True)
    doc_part = doc.part

    existing = None
    for rel in doc_part.rels.values():
        if 'comments' in rel.reltype:
            existing = rel
            break

    if existing:
        existing.target_part._blob = comments_bytes
    else:
        comments_part = Part(
            PackURI('/word/comments.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
            comments_bytes,
            doc_part.package
        )
        doc_part.relate_to(comments_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')

def add_comment(doc, comments_element, search_text, comment_text, location_hint=None):
    """Search in paragraphs and table cells, then add comment."""
    comment_id = comment_id_counter[0]
    comment_id_counter[0] += 1

    # Search paragraphs
    for para in doc.paragraphs:
        if search_text.lower() in (para.text or "").lower():
            add_comment_to_xml(comments_element, comment_id, AUTHOR, DATE_STR, comment_text)
            add_comment_markers(para, search_text, comment_id)
            return comment_id, True

    # Search table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if search_text.lower() in (para.text or "").lower():
                        add_comment_to_xml(comments_element, comment_id, AUTHOR, DATE_STR, comment_text)
                        add_comment_markers(para, search_text, comment_id)
                        return comment_id, True

    return comment_id, False

def main():
    doc = Document(DOC_PATH)

    findings = [
        # (search_text, comment_text, category)
        ("文章编号:", "文章编号为空，请补充文章编号。", "格式规范"),
        ("Using a representative mine ecological restoration area as an empirical case, this study clarifies how mine restoration reshapes the spatiotemporal dynamics of ecosystem services",
         "英文摘要内容明显短于中文摘要，且似乎未完整对应中文摘要的四部分结构（目的、方法、结果、结论）。建议对照中文摘要补充完整。", "摘要"),
        ("Gaozhi Peng",
         "英文作者姓名格式不规范。按照中文期刊规范，应为姓前名后，姓全拼，名缩写。建议改为 Gao Z P 或 Gao Zhipeng。", "作者信息"),
        ("a case study of jungar banner",
         "英文题名中地名拼写为 Jungar Banner，但正文中统一使用 Zhunge′er Banner，请统一地名拼写。", "题名"),
        ("式中:为流域内每个栅格单元的年产水量",
         "公式注解中变量符号缺失（“式中：”后直接接“为”）。公式变量可能以公式对象形式插入但未在注解中显示。请检查并补全每个变量的符号说明。", "公式"),
        ("式中:为总碳储量(t/hm2)",
         "公式注解中变量符号缺失（“式中：”后直接接“为”）。请补全变量符号。", "公式"),
        ("式中:为栅格x的土壤保持量(t)",
         "公式注解中变量符号缺失（“式中：”后直接接“为”）。请补全变量符号。", "公式"),
        ("式中:为土地利用i类中x栅格的生境质量",
         "公式注解中变量符号缺失（“式中：”后直接接“为”）。请补全变量符号。", "公式"),
        ("式中:为皮尔逊相关系数",
         "公式注解中变量符号缺失（“式中：”后直接接“为”）。请补全变量符号。", "公式"),
        ("式中:为第𝑖个空间单元中的因变量",
         "公式注解中变量符号缺失（“式中：”后直接接“为”）。请补全变量符号。", "公式"),
        ("4.讨 论",
         "节标题“4.讨 论”中间有空格，与其他节标题（如“3.结果与分析”）格式不一致。建议去掉空格，改为“4.讨论”。", "格式规范"),
        ("5.结 论",
         "节标题“5.结 论”中间有空格，与其他节标题格式不一致。建议去掉空格，改为“5.结论”。", "格式规范"),
        ("Acta Ecologica Sinica, 2023(16):1-14",
         "英文参考文献卷号缺失。应为 Acta Ecologica Sinica, 2023, 43(16):1-14。", "参考文献"),
        ("358-36",
         "英文参考文献页码不完整，疑似缺漏。请核实并补全。", "参考文献"),
        ("China mining Magazine",
         "期刊名大小写不规范。建议改为 China Mining Magazine。", "参考文献"),
        ("Fig. 3Spatial distribution",
         "图题中英文缩写 Fig. 3 与正文之间缺少空格，与其他图题格式不一致。建议改为 Fig. 3 Spatial distribution。", "图表格式"),
        ("四项具有代表性的生态系统服务",
         "前文使用汉字数字“四项”，后文出现“4类生态系统服务”（阿拉伯数字）。全文数字用法请保持一致，建议统一使用阿拉伯数字。", "文字规范"),
        ("https:∥",
         "表格中URL使用了特殊字符“∥”，应使用半角斜杠“//”。请检查并修正所有网址格式。", "图表格式"),
    ]

    inserted = []
    not_found = []

    for search_text, comment_text, category in findings:
        cid, ok = add_comment(doc, comments_element, search_text, comment_text)
        if ok:
            inserted.append((search_text[:50], comment_text, category))
            print(f"[OK] Comment {cid}: {category}")
        else:
            not_found.append((search_text, comment_text, category))
            print(f"[NOT FOUND] Comment {cid}: {search_text[:60]}")

    attach_comments_part(doc, comments_element)
    doc.save(DOC_PATH)

    print(f"\nInserted {len(inserted)} comments.")
    if not_found:
        print(f"Not found ({len(not_found)}):")
        for s, c, cat in not_found:
            print(f"  - [{cat}] {s[:60]}")

    # Summary by category
    from collections import defaultdict
    cats = defaultdict(list)
    for s, c, cat in inserted:
        cats[cat].append(c)
    print("\n--- Summary ---")
    for cat, items in cats.items():
        print(f"{cat}: {len(items)} comments")

if __name__ == "__main__":
    main()
